# views.py

import json
from operator import itemgetter
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from django.views.generic import ListView, CreateView, UpdateView, DeleteView, DetailView
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.auth.decorators import login_required, permission_required # Импортируем permission_required
from django.utils import timezone
from django.http import HttpResponse
from django.db.models import Count, Q
from itertools import groupby
from django.core.exceptions import PermissionDenied
from django.views.decorators.http import require_POST
from django.db import transaction

from docx import Document
from docxtpl import DocxTemplate

from .models import Applicant, Order, Student 
from .forms import ApplicantForm, OrderForm
from .filters import ApplicantFilter
import pandas as pd
from django.contrib import messages
from .models import Specialty, Qualification
import difflib
from django.core.serializers.json import DjangoJSONEncoder

@require_POST # Эта функция будет принимать только POST-запросы
@permission_required('enrollment.change_applicant', raise_exception=True)
def toggle_enrollment_status(request, pk):
    """Переключает статус абитуриента 'К зачислению'."""
    applicant = get_object_or_404(Applicant, pk=pk)
    applicant.is_ready_for_enrollment = not applicant.is_ready_for_enrollment
    applicant.save()
    return redirect('applicant_list')

@login_required
def applicant_list(request):
    """Список абитуриентов с фильтрацией."""
    f = ApplicantFilter(request.GET, queryset=Applicant.objects.all())
    return render(request, 'enrollment/applicant_list.html', {'filter': f})

# ИСПРАВЛЕНО: Заменяем @writer_required на стандартный декоратор
@permission_required('enrollment.add_applicant', raise_exception=True)
def add_applicant(request):
    if request.method == 'POST':
        form = ApplicantForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('applicant_list')
    else:
        form = ApplicantForm()
    return render(request, 'enrollment/applicant_form.html', {'form': form})

# ИСПРАвлено: Заменяем @editor_required на стандартный декоратор
@permission_required('enrollment.change_applicant', raise_exception=True)
def edit_applicant(request, pk):
    applicant = get_object_or_404(Applicant, pk=pk)
    if request.method == 'POST':
        form = ApplicantForm(request.POST, instance=applicant)
        if form.is_valid():
            form.save()
            return redirect('applicant_list')
    else:
        form = ApplicantForm(instance=applicant)
    return render(request, 'enrollment/applicant_form.html', {'form': form})


# --- Представления для Приказов (логика уже исправлена) ---
@permission_required('enrollment.add_order', raise_exception=True)
def order_create(request):
    """Создание нового приказа с фильтрацией абитуриентов."""
    # Используем тот же фильтр, что и на главной
    applicant_filter = ApplicantFilter(request.GET, queryset=Applicant.objects.filter(is_ready_for_enrollment=True))
    
    if request.method == 'POST':
        form = OrderForm(request.POST)
        if form.is_valid():
            # Получаем список ID абитуриентов, выбранных в таблице
            selected_applicant_ids = request.POST.getlist('selected_applicants')
            
            if not selected_applicant_ids:
                messages.error(request, "Не выбрано ни одного абитуриента!")
                # Возвращаемся на ту же страницу, сохраняя фильтры
                return render(request, 'enrollment/order_form.html', {
                    'form': form,
                    'filter': applicant_filter,
                })

            # Создаем приказ, но пока не сохраняем в базу
            order = form.save(commit=False)
            order.created_by = request.user
            order.save() # Сохраняем, чтобы получить ID

            # Добавляем выбранных абитуриентов в приказ
            order.applicants.set(selected_applicant_ids)
            
            messages.success(request, "Черновик приказа успешно создан.")
            return redirect('order_detail', pk=order.pk)
    else:
        form = OrderForm()

    return render(request, 'enrollment/order_form.html', {
        'form': form,
        'filter': applicant_filter,
    })


@permission_required('enrollment.change_order', raise_exception=True)
def order_update(request, pk):
    """Редактирование приказа с фильтрацией."""
    order = get_object_or_404(Order, pk=pk)
    
    # Проверка прав доступа, как в старом test_func
    if not (request.user.is_superuser or order.status == 'draft'):
        raise PermissionDenied("Вы можете редактировать только черновики.")

    applicant_filter = ApplicantFilter(request.GET, queryset=Applicant.objects.filter(is_ready_for_enrollment=True))
    
    if request.method == 'POST':
        form = OrderForm(request.POST, instance=order)
        if form.is_valid():
            selected_applicant_ids = request.POST.getlist('selected_applicants')
            
            if not selected_applicant_ids:
                messages.error(request, "Не выбрано ни одного абитуриента!")
                return render(request, 'enrollment/order_form.html', {
                    'form': form,
                    'filter': applicant_filter,
                    'object': order, # для шаблона
                })

            order = form.save()
            order.applicants.set(selected_applicant_ids)
            
            messages.success(request, "Приказ успешно обновлен.")
            return redirect('order_detail', pk=order.pk)
    else:
        form = OrderForm(instance=order)

    return render(request, 'enrollment/order_form.html', {
        'form': form,
        'filter': applicant_filter,
        'object': order, # для шаблона, чтобы знать ID уже добавленных
    })

class OrderListView(LoginRequiredMixin, ListView):
    model = Order
    template_name = 'enrollment/order_list.html'
    context_object_name = 'orders'

class OrderDetailView(LoginRequiredMixin, DetailView):
    model = Order
    template_name = 'enrollment/order_detail.html'

class OrderDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Order
    template_name = 'enrollment/order_confirm_delete.html'
    success_url = reverse_lazy('order_list')

    def test_func(self):
        user = self.request.user
        if not user.has_perm('enrollment.delete_order'):
            return False
        order = self.get_object()
        return user.is_superuser or order.status == 'draft'

@require_POST
@permission_required('enrollment.change_order', raise_exception=True)
@transaction.atomic # Гарантирует, что все операции либо выполнятся успешно, либо отменятся
def sign_order(request, pk):
    """
    Подписывает приказ и выполняет зачисление абитуриентов,
    перенося их в таблицу студентов.
    """
    order = get_object_or_404(Order, pk=pk)
    order_number_from_form = request.POST.get('order_number')

    if not order_number_from_form:
        messages.error(request, "Номер приказа не может быть пустым.")
        return redirect('order_detail', pk=pk)

    if Order.objects.filter(order_number=order_number_from_form).exclude(pk=order.pk).exists():
        messages.error(request, f"Ошибка: Приказ с номером '{order_number_from_form}' уже существует!")
        return redirect('order_detail', pk=pk)

    if order.status == 'draft':
        # 1. Подписываем сам приказ
        order.status = 'signed'
        order.order_number = order_number_from_form
        order.order_date = timezone.now().date()
        order.signed_by = request.user
        order.save()

        # 2. Начинаем процесс перевода абитуриентов в студенты
        applicants_to_enroll = list(order.applicants.all()) # Создаем копию списка
        for applicant in applicants_to_enroll:
            # Создаем студента, копируя все поля
            student = Student.objects.create(
                full_name=applicant.full_name,
                iin=applicant.iin,
                specialty=applicant.specialty,
                qualification=applicant.qualification,
                study_form=applicant.study_form,
                base_education=applicant.base_education,
                school=applicant.school,
                graduation_year=applicant.graduation_year,
                with_honors=applicant.with_honors,
                gpa=applicant.gpa,
                birth_date=applicant.birth_date,
                study_language=applicant.study_language,
                social_status=applicant.social_status,
                citizenship=applicant.citizenship,
                nationality=applicant.nationality,
                gender=applicant.gender,
                phone_number=applicant.phone_number,
                home_address=applicant.home_address,
                parents_info=applicant.parents_info,
                has_incomplete_docs=applicant.has_incomplete_docs,
                needs_dormitory=applicant.needs_dormitory,
                payment_type=applicant.payment_type,
                application_date=applicant.application_date,
                photo=applicant.photo
            )
            # Добавляем приказ в историю перемещений студента
            student.movement_history.add(order)
            
            # Удаляем абитуриента
            applicant.delete()

        messages.success(request, f"Приказ №{order.order_number} подписан. Зачислено студентов: {len(applicants_to_enroll)}.")
    else:
        messages.warning(request, "Этот приказ уже был подписан ранее.")

    return redirect('order_detail', pk=pk)


@login_required
def generate_order_docx(request, pk):
    order = get_object_or_404(Order, pk=pk)
    document = Document()
    document.add_heading(f"ПРИКАЗ № {order.order_number or 'Б/Н'}", level=1)
    if order.order_date:
        document.add_heading(f"от {order.order_date.strftime('%d.%m.%Y')}", level=2)
    document.add_paragraph(f"О зачислении в группу: {order.group_name}")
    
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'ФИО'
    hdr_cells[2].text = 'Тип оплаты'

    for i, applicant in enumerate(order.applicants.all(), 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = applicant.full_name
        row_cells[2].text = applicant.payment_type
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="prikaz_{order.pk}.docx"'
    document.save(response)
    
    return response

@login_required
def export_applicants_to_excel(request):
    """
    Экспортирует данные абитуриентов в Excel файл с учетом новых полей.
    """
    f = ApplicantFilter(request.GET, queryset=Applicant.objects.select_related('specialty', 'qualification').all())
    queryset = f.qs

    data = list(queryset.values(
        'full_name', 'iin', 'specialty__name', 'qualification__name', 
        'payment_type', 'payment_status', 'study_form', 'base_education', 
        'gpa', 'birth_date', 'school', 'graduation_year', 'with_honors', 
        'study_language', 'social_status', 'citizenship', 'nationality', 
        'gender', 'phone_number', 'home_address', 'parents_info', 
        'has_incomplete_docs', 'needs_dormitory', 'application_date', 
        'is_ready_for_enrollment'
    ))
    df = pd.DataFrame(data)

    # Переименовываем столбцы для красоты
    df.rename(columns={
        'full_name': 'ФИО', 'iin': 'ИИН', 'specialty__name': 'Специальность', 
        'qualification__name': 'Квалификация', 'payment_type': 'Тип финансирования',
        'payment_status': 'Статус оплаты', 'study_form': 'Форма обучения',
        'base_education': 'На базе', 'gpa': 'Средний балл аттестата', 
        'birth_date': 'Дата рождения', 'school': 'Школа/Колледж',
        'graduation_year': 'Год окончания', 'with_honors': 'С отличием',
        'study_language': 'Язык обучения', 'social_status': 'СП',
        'citizenship': 'Гражданство', 'nationality': 'Национальность',
        'gender': 'Пол', 'phone_number': 'Контакты', 'home_address': 'Дом. адрес',
        'parents_info': 'ФИО родителей', 'has_incomplete_docs': 'Неполный пакет документов',
        'needs_dormitory': 'Общежитие', 'application_date': 'Дата подачи заявления',
        'is_ready_for_enrollment': 'Готов к зачислению'
    }, inplace=True)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="applicants_{timezone.now().strftime("%Y-%m-%d")}.xlsx"'
    df.to_excel(response, index=False)
    return response


@permission_required('enrollment.add_applicant', raise_exception=True)
def import_applicants_from_excel(request):
    """
    Импортирует абитуриентов из Excel файла (исправленная и обновленная версия).
    """
    if request.method == 'POST':
        excel_file = request.FILES.get('excel_file')
        if not excel_file:
            messages.error(request, "Файл не был предоставлен.")
            return redirect('applicant_list')

        try:
            # Читаем все данные как строки, чтобы избежать проблем с форматами
            df = pd.read_excel(excel_file, dtype=str).where(pd.notna, None)
            
            required_columns = ['ФИО', 'ИИН', 'Специальность', 'Квалификация']
            if not all(col in df.columns for col in required_columns):
                messages.error(request, f"Ошибка: в файле отсутствуют обязательные столбцы. Требуются: {', '.join(required_columns)}")
                return redirect('applicant_list')

            # Кэшируем данные, чтобы не делать запросы в цикле
            all_specialties = {s.name: s for s in Specialty.objects.all()}
            all_qualifications = {f"{q.name}|{q.specialty.name}": q for q in Qualification.objects.select_related('specialty').all()}
            
            imported_count = 0
            for index, row in df.iterrows():
                iin = row.get('ИИН')
                spec_name = row.get('Специальность')
                qual_name = row.get('Квалификация')
                
                if not all([iin, spec_name, qual_name, row.get('ФИО')]):
                    messages.warning(request, f"Строка {index + 2}: пропущена из-за отсутствия ИИН, ФИО, специальности или квалификации.")
                    continue

                specialty = all_specialties.get(spec_name.strip())
                if not specialty:
                    messages.warning(request, f"Строка {index + 2} ({row['ФИО']}): специальность '{spec_name}' не найдена.")
                    continue
                
                qual_key = f"{qual_name.strip()}|{specialty.name}"
                qualification = all_qualifications.get(qual_key)
                if not qualification:
                    messages.warning(request, f"Строка {index + 2} ({row['ФИО']}): квалификация '{qual_name}' не найдена для специальности '{specialty.name}'.")
                    continue

                # Обрабатываем дату
                try:
                    app_date = pd.to_datetime(row.get('Дата подачи заявления')).date() if row.get('Дата подачи заявления') else timezone.now().date()
                    birth_date = pd.to_datetime(row.get('Дата рождения')).date() if row.get('Дата рождения') else None
                except (ValueError, TypeError):
                    app_date = timezone.now().date()
                    birth_date = None
                    messages.warning(request, f"Строка {index + 2} ({row['ФИО']}): неверный формат даты. Даты не были установлены.")

                Applicant.objects.update_or_create(
                    iin=iin,
                    defaults={
                        'full_name': row.get('ФИО'),
                        'specialty': specialty,
                        'qualification': qualification,
                        'application_date': app_date,
                        'birth_date': birth_date,
                        # ... добавьте сюда остальные поля из Excel по аналогии ...
                        'gpa': row.get('Средний балл аттестата'),
                        'study_language': row.get('Язык обучения'),
                        'payment_type': row.get('Тип финансирования'),
                    }
                )
                imported_count += 1

            messages.success(request, f"Импорт успешно завершен! Обработано записей: {imported_count}.")
        
        except Exception as e:
            messages.error(request, f"Произошла критическая ошибка при импорте: {e}")

    return redirect('applicant_list')


@login_required
def dashboard_view(request):
    """
    Отображает дашборд со сводной статистикой по абитуриентам.
    """
    qualifications_data = Qualification.objects.annotate(
        total_applicants=Count('applicant'),

        # Подсчет по базе 9
        base9_rus=Count('applicant', filter=Q(
            applicant__base_education='9 классов',
            applicant__study_language__icontains='рус'
        ) & ~Q(applicant__study_form='Дуальная')),
        base9_kaz=Count('applicant', filter=Q(
            applicant__base_education='9 классов',
            applicant__study_language__icontains='каз'
        ) & ~Q(applicant__study_form='Дуальная')),

        # Подсчет по базе 11
        base11_rus=Count('applicant', filter=Q(
            applicant__base_education='11 классов',
            applicant__study_language__icontains='рус'
        ) & ~Q(applicant__study_form='Дуальная')),
        base11_kaz=Count('applicant', filter=Q(
            applicant__base_education='11 классов',
            applicant__study_language__icontains='каз'
        ) & ~Q(applicant__study_form='Дуальная')),

        # ДОБАВЛЕНО: Подсчет по базе ТиПО
        tipo_rus=Count('applicant', filter=Q(
            applicant__base_education='ТиПО',
            applicant__study_language__icontains='рус'
        ) & ~Q(applicant__study_form='Дуальная')),
        tipo_kaz=Count('applicant', filter=Q(
            applicant__base_education='ТиПО',
            applicant__study_language__icontains='каз'
        ) & ~Q(applicant__study_form='Дуальная')),

        # Подсчет дуальщиков
        dual=Count('applicant', filter=Q(applicant__study_form='Дуальная')),

    ).values(
        'name', 'specialty__name', 'specialty_id', 'total_applicants',
        'base9_rus', 'base9_kaz', 'base11_rus', 'base11_kaz',
        'tipo_rus', 'tipo_kaz',  # ДОБАВЛЕНО
        'dual'
    ).order_by('specialty__name', 'name')

    # --- Группировка остается без изменений ---
    grouped_data = []
    for key, group in groupby(qualifications_data, key=lambda x: x['specialty__name']):
        qualifications = list(group)
        specialty_total = sum(q['total_applicants'] for q in qualifications)
        grouped_data.append({
            'specialty_name': key,
            'qualifications': qualifications,
            'specialty_total': specialty_total,
        })

    # --- Обновляем итоговые суммы ---
    totals = {
        'total_base9_rus': sum(item['base9_rus'] for item in qualifications_data),
        'total_base9_kaz': sum(item['base9_kaz'] for item in qualifications_data),
        'total_base11_rus': sum(item['base11_rus'] for item in qualifications_data),
        'total_base11_kaz': sum(item['base11_kaz'] for item in qualifications_data),
        'total_tipo_rus': sum(item['tipo_rus'] for item in qualifications_data),  # ДОБАВЛЕНО
        'total_tipo_kaz': sum(item['tipo_kaz'] for item in qualifications_data),  # ДОБАВЛЕНО
        'total_dual': sum(item['dual'] for item in qualifications_data),
        'grand_total': sum(item['specialty_total'] for item in grouped_data)
    }

    context = {
        'grouped_data': grouped_data,
        'totals': totals,
    }
     # --- НАЧАЛО: НОВЫЙ КОД ДЛЯ ГРАФИКОВ ---

    # 1. Данные для круговой диаграммы (по базе)
    base_data = Applicant.objects.values('base_education', 'study_form').annotate(count=Count('id'))
    
    chart_base_counts = {
        '9 классов': 0,
        '11 классов': 0,
        'ТиПО': 0,
        'Дуальное': 0,
    }
    for item in base_data:
        if item['study_form'] == 'Дуальная':
            chart_base_counts['Дуальное'] += item['count']
        else:
            if item['base_education'] in chart_base_counts:
                chart_base_counts[item['base_education']] += item['count']

    # 2. Данные для гистограммы (по специальностям)
    chart_specialty_labels = [data['specialty_name'] for data in grouped_data]
    chart_specialty_counts = [data['specialty_total'] for data in grouped_data]

    # --- КОНЕЦ: НОВЫЙ КОД ДЛЯ ГРАФИКОВ ---

    context = {
        'grouped_data': grouped_data,
        'totals': totals,
        # Передаем данные для графиков в контекст
        'chart_base_labels': json.dumps(list(chart_base_counts.keys())),
        'chart_base_data': json.dumps(list(chart_base_counts.values())),
        'chart_specialty_labels': json.dumps(chart_specialty_labels, cls=DjangoJSONEncoder),
        'chart_specialty_data': json.dumps(chart_specialty_counts),
    }
    return render(request, 'enrollment/dashboard.html', context)

# НОВАЯ ФУНКЦИЯ ДЛЯ ЭКСПОРТА
@login_required
def export_dashboard_excel(request):
    """
    Формирует и отдает Excel-файл на основе данных из таблицы дашборда.
    """
    # Этот код полностью повторяет логику получения данных из dashboard_view.
    # В идеале, эту логику можно вынести в отдельную функцию, чтобы избежать дублирования.
    qualifications_data = Qualification.objects.annotate(
        total_applicants=Count('applicant'),
        base9_rus=Count('applicant', filter=Q(applicant__base_education='9 классов', applicant__study_language__icontains='рус') & ~Q(applicant__study_form='Дуальная')),
        base9_kaz=Count('applicant', filter=Q(applicant__base_education='9 классов', applicant__study_language__icontains='каз') & ~Q(applicant__study_form='Дуальная')),
        base11_rus=Count('applicant', filter=Q(applicant__base_education='11 классов', applicant__study_language__icontains='рус') & ~Q(applicant__study_form='Дуальная')),
        base11_kaz=Count('applicant', filter=Q(applicant__base_education='11 классов', applicant__study_language__icontains='каз') & ~Q(applicant__study_form='Дуальная')),
        tipo_rus=Count('applicant', filter=Q(applicant__base_education='ТиПО', applicant__study_language__icontains='рус') & ~Q(applicant__study_form='Дуальная')),
        tipo_kaz=Count('applicant', filter=Q(applicant__base_education='ТиПО', applicant__study_language__icontains='каз') & ~Q(applicant__study_form='Дуальная')),
        dual=Count('applicant', filter=Q(applicant__study_form='Дуальная')),
    ).values(
        'name', 'specialty__name', 'total_applicants', 'base9_rus', 'base9_kaz',
        'base11_rus', 'base11_kaz', 'tipo_rus', 'tipo_kaz', 'dual'
    ).order_by('specialty__name', 'name')

    # Формируем данные для DataFrame
    records = []
    for q in qualifications_data:
        records.append({
            'Образовательная программа': q['specialty__name'],
            'Квалификация': q['name'],
            'База 9 (рус)': q['base9_rus'],
            'База 9 (каз)': q['base9_kaz'],
            'База 11 (рус)': q['base11_rus'],
            'База 11 (каз)': q['base11_kaz'],
            'База ТиПО (рус)': q['tipo_rus'],
            'База ТиПО (каз)': q['tipo_kaz'],
            'Дуальное': q['dual'],
            'Всего по квалификации': q['total_applicants']
        })
    df = pd.DataFrame(records)

    # Создаем HTTP-ответ с Excel файлом
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="dashboard_report_{timezone.now().strftime("%Y-%m-%d")}.xlsx"'
    df.to_excel(response, index=False)
    
    return response

@login_required
def generate_order_docx(request, pk):
    order = get_object_or_404(Order, pk=pk)
    doc = DocxTemplate("templates/enrollment/docx_templates/prikaz_template.docx")

    # Получаем всех студентов и сортируем их в правильном порядке для группировки
    all_students = order.student_set.select_related(
        'specialty', 'qualification'
    ).order_by(
        'qualification__is_worker_qualification', # Сначала обычные, потом рабочие
        'specialty__name', 
        'qualification__name', 
        'base_education', 
        'study_language', 
        'full_name'
    )

    # Разделяем студентов на две группы
    regular_students = [s for s in all_students if not s.qualification.is_worker_qualification]
    worker_students = [s for s in all_students if s.qualification.is_worker_qualification]
    
    # Функция для группировки списка студентов
    def group_student_list(student_list):
        grouped_data = []
        # Группируем по специальности
        for specialty, specialty_group in groupby(student_list, key=lambda s: s.specialty):
            sub_groups = []
            
            # ИСПРАВЛЕНИЕ: Используем lambda вместо itemgetter
            sub_group_key = lambda s: (s.qualification, s.base_education, s.study_language)
            
            # Группируем внутренний список по составному ключу
            for key, group in groupby(specialty_group, key=sub_group_key):
                sub_groups.append({
                    'qualification': key[0],
                    'base_education': key[1],
                    'study_language': key[2],
                    'students': list(group)
                })
            grouped_data.append({'specialty': specialty, 'sub_groups': sub_groups})
        return grouped_data

    # Готовим контекст для шаблона
    months = ["января", "февраля", "марта", "апреля", "мая", "июня", "июля", "августа", "сентября", "октября", "ноября", "декабря"]
    context = {
        'order_number': order.order_number,
        'order_date': order.order_date,
        'order_date_month_str': months[order.order_date.month - 1] if order.order_date else '',
        'title': order.title,
        'preamble': order.preamble,
        'grouped_regular_students': group_student_list(regular_students),
        'grouped_worker_students': group_student_list(worker_students),
        'signer_name': order.signed_by.get_full_name() if order.signed_by else " ",
        'signer_title': order.signer_title,
    }
    
    doc.render(context)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="prikaz_{order.order_number}.docx"'
    doc.save(response)
    
    return response